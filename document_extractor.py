"""
Document extraction module for Texas Child Care Solutions scraper
Handles text extraction from .docx and .xlsx files
"""

import os
import logging
import requests
import hashlib
from urllib.parse import urlparse
from typing import Optional, Dict, Any

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logging.warning("python-docx not available. DOCX extraction will be disabled.")

try:
    from openpyxl import load_workbook
    OPENPYXL_AVAILABLE = True
except ImportError:
    OPENPYXL_AVAILABLE = False
    logging.warning("openpyxl not available. XLSX extraction will be disabled.")

import config

# Set up logging
logger = logging.getLogger(__name__)


class DocumentExtractor:
    """Handles extraction from .docx and .xlsx documents."""

    def __init__(self, output_dir: str = None):
        """
        Initialize document extractor.

        Args:
            output_dir: Directory to save downloaded documents
        """
        self.output_dir = output_dir or config.PAGES_DIR
        self.user_agent = config.USER_AGENT
        self.max_size_bytes = getattr(config, 'MAX_DOCUMENT_SIZE_MB', 50) * 1024 * 1024
        self.timeout = config.TIMEOUT_PER_PAGE

        # Create output directory if it doesn't exist
        os.makedirs(self.output_dir, exist_ok=True)

    def _generate_doc_id(self, url: str) -> str:
        """Generate a unique ID for a document URL."""
        return hashlib.md5(url.encode()).hexdigest()

    def _get_filename_from_url(self, url: str) -> str:
        """Extract filename from URL."""
        parsed = urlparse(url)
        filename = os.path.basename(parsed.path)
        if not filename:
            filename = f"{self._generate_doc_id(url)}.tmp"
        return filename

    def download_document(self, url: str) -> Optional[str]:
        """
        Download document from URL.

        Args:
            url: URL of the document

        Returns:
            Path to downloaded document, or None if download failed
        """
        try:
            logger.info(f"Downloading document: {url}")

            # Make HEAD request to check size
            headers = {'User-Agent': self.user_agent}
            head_response = requests.head(url, headers=headers, timeout=10, allow_redirects=True)

            # Check file size
            content_length = head_response.headers.get('Content-Length')
            if content_length:
                size_mb = int(content_length) / (1024 * 1024)
                if int(content_length) > self.max_size_bytes:
                    logger.warning(f"Document too large ({size_mb:.1f} MB): {url}")
                    return None
                logger.debug(f"Document size: {size_mb:.1f} MB")

            # Download the document
            response = requests.get(url, headers=headers, timeout=self.timeout, stream=True)
            response.raise_for_status()

            # Save to file
            filename = self._get_filename_from_url(url)
            filepath = os.path.join(self.output_dir, filename)

            with open(filepath, 'wb') as f:
                for chunk in response.iter_content(chunk_size=8192):
                    f.write(chunk)

            logger.debug(f"Downloaded document to: {filepath}")
            return filepath

        except requests.RequestException as e:
            logger.error(f"Failed to download document {url}: {e}")
            return None
        except Exception as e:
            logger.error(f"Unexpected error downloading document {url}: {e}")
            return None

    def extract_text_from_docx(self, docx_path: str) -> Optional[str]:
        """
        Extract text from DOCX file.

        Args:
            docx_path: Path to DOCX file

        Returns:
            Extracted text, or None if extraction failed
        """
        if not DOCX_AVAILABLE:
            logger.error("python-docx not available. Cannot extract text.")
            return None

        try:
            logger.info(f"Extracting text from DOCX: {docx_path}")

            doc = Document(docx_path)
            text_parts = []

            # Extract text from paragraphs
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    text_parts.append(text)

            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text:
                            row_text.append(cell_text)
                    if row_text:
                        text_parts.append(' | '.join(row_text))

            full_text = '\n\n'.join(text_parts)
            word_count = len(full_text.split())

            logger.info(f"Extracted {word_count} words from DOCX")

            return full_text.strip()

        except Exception as e:
            logger.error(f"Failed to extract text from {docx_path}: {e}")
            return None

    def extract_text_from_xlsx(self, xlsx_path: str) -> Optional[str]:
        """
        Extract text from XLSX file.

        Args:
            xlsx_path: Path to XLSX file

        Returns:
            Extracted text, or None if extraction failed
        """
        if not OPENPYXL_AVAILABLE:
            logger.error("openpyxl not available. Cannot extract text.")
            return None

        try:
            logger.info(f"Extracting text from XLSX: {xlsx_path}")

            workbook = load_workbook(xlsx_path, read_only=True, data_only=True)
            text_parts = []

            for sheet_name in workbook.sheetnames:
                sheet = workbook[sheet_name]

                # Add sheet name as heading
                text_parts.append(f"=== Sheet: {sheet_name} ===")

                # Extract cell values
                for row in sheet.iter_rows(values_only=True):
                    # Filter out None values and convert to strings
                    row_values = [str(cell).strip() for cell in row if cell is not None and str(cell).strip()]
                    if row_values:
                        text_parts.append(' | '.join(row_values))

            workbook.close()

            full_text = '\n'.join(text_parts)
            word_count = len(full_text.split())

            logger.info(f"Extracted {word_count} words from XLSX ({len(workbook.sheetnames)} sheets)")

            return full_text.strip()

        except Exception as e:
            logger.error(f"Failed to extract text from {xlsx_path}: {e}")
            return None

    def is_document_url(self, url: str) -> bool:
        """
        Check if URL appears to be a document (.docx or .xlsx).

        Args:
            url: URL to check

        Returns:
            True if URL appears to be a document
        """
        url_lower = url.lower()
        return url_lower.endswith('.docx') or url_lower.endswith('.xlsx')

    def process_document_url(self, url: str, metadata: Dict[str, Any] = None) -> Optional[Dict[str, Any]]:
        """
        Download and extract text from document URL.

        Args:
            url: URL of the document
            metadata: Optional metadata to include

        Returns:
            Dictionary with extracted content and metadata, or None if failed
        """
        doc_id = self._generate_doc_id(url)

        # Download document
        doc_path = self.download_document(url)
        if not doc_path:
            return None

        # Determine document type and extract text
        text = None
        if url.lower().endswith('.docx'):
            text = self.extract_text_from_docx(doc_path)
            content_type = 'docx'
        elif url.lower().endswith('.xlsx'):
            text = self.extract_text_from_xlsx(doc_path)
            content_type = 'xlsx'
        else:
            logger.warning(f"Unknown document type: {url}")
            return None

        if not text:
            return None

        # Build result
        result = {
            'doc_id': doc_id,
            'url': url,
            'filename': os.path.basename(doc_path),
            'text': text,
            'word_count': len(text.split()),
            'content_type': content_type,
        }

        # Add additional metadata if provided
        if metadata:
            result.update(metadata)

        # Clean up downloaded file (optional)
        # os.remove(doc_path)

        return result


# Convenience function for quick extraction
def extract_document(url: str, output_dir: str = None) -> Optional[Dict[str, Any]]:
    """
    Quick function to extract text from a document URL.

    Args:
        url: URL of the document
        output_dir: Directory to save extracted text

    Returns:
        Dictionary with extracted content and metadata
    """
    extractor = DocumentExtractor(output_dir)
    return extractor.process_document_url(url)


if __name__ == '__main__':
    # Test the extractor
    logging.basicConfig(level=logging.INFO, format=config.LOG_FORMAT)

    # Test URLs (replace with actual document URLs if needed)
    test_docx_url = "https://www.twc.texas.gov/sites/default/files/ccel/docs/ccms.docx"
    test_xlsx_url = "https://www.twc.texas.gov/sites/default/files/ccel/docs/child-care-numbers-data-monthly.xlsx"

    extractor = DocumentExtractor()

    print("\n=== Testing DOCX extraction ===")
    result = extractor.process_document_url(test_docx_url)
    if result:
        print(f"Successfully extracted {result['word_count']} words from DOCX")
        print(f"First 200 characters: {result['text'][:200]}")
    else:
        print("DOCX extraction failed")

    print("\n=== Testing XLSX extraction ===")
    result = extractor.process_document_url(test_xlsx_url)
    if result:
        print(f"Successfully extracted {result['word_count']} words from XLSX")
        print(f"First 200 characters: {result['text'][:200]}")
    else:
        print("XLSX extraction failed")
